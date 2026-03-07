"""Offline text extraction utilities for PRIVGUARD AI.

Supports plain text-like files, DOCX documents, PDF extraction, and image OCR
using local Tesseract. Sharp preprocessing improves OCR for unclear photos.
"""

from __future__ import annotations

import io
import json
import logging
import os
import shutil
from contextlib import redirect_stderr, redirect_stdout
from pathlib import Path
from typing import Any, Dict, List

from PIL import Image, ImageFilter, ImageOps
import pytesseract

try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency import path
    Document = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - optional dependency import path
    PdfReader = None

try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover - optional dependency import path
    fitz = None

try:
    import cv2
    import numpy as np

    _CV2_AVAILABLE = True
except ImportError:
    _CV2_AVAILABLE = False

try:
    from config_loader import load_system_config
except ImportError:
    load_system_config = None

TEXT_SUFFIXES = {".txt", ".md", ".csv", ".log"}
DOCX_SUFFIXES = {".docx"}
PDF_SUFFIXES = {".pdf"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp"}

_DEFAULT_OCR_CONFIG: Dict[str, Any] = {
    "sharp_preprocessing": True,
    "min_text_height_px": 1200,
    "max_dimension_px": 4000,
    "pdf_scale": 3,
}

_OCR_OVERRIDE_PATH = Path("instance/ocr_override.json")


def _get_ocr_config() -> Dict[str, Any]:
    """Return OCR settings from system config and optional dashboard override."""
    out = dict(_DEFAULT_OCR_CONFIG)
    if load_system_config:
        try:
            cfg = load_system_config().get("ocr") or {}
            for key in out:
                if key in cfg and cfg[key] is not None:
                    out[key] = cfg[key]
        except Exception:
            pass

    if _OCR_OVERRIDE_PATH.exists():
        try:
            raw = _OCR_OVERRIDE_PATH.read_text(encoding="utf-8")
            override = json.loads(raw)
            if "sharp_preprocessing" in override:
                out["sharp_preprocessing"] = bool(override["sharp_preprocessing"])
        except Exception:
            pass
    return out


def get_ocr_config() -> Dict[str, Any]:
    """Public API for current OCR settings."""
    return _get_ocr_config()


def set_ocr_override(sharp_preprocessing: bool) -> None:
    """Persist Sharp OCR toggle from dashboard."""
    _OCR_OVERRIDE_PATH.parent.mkdir(parents=True, exist_ok=True)
    _OCR_OVERRIDE_PATH.write_text(
        json.dumps({"sharp_preprocessing": sharp_preprocessing}, indent=2),
        encoding="utf-8",
    )


def _configure_tesseract_cmd() -> None:
    """Configure Tesseract path for Windows-friendly local setups."""
    env_cmd = os.environ.get("TESSERACT_CMD", "").strip()
    if env_cmd and Path(env_cmd).exists():
        pytesseract.pytesseract.tesseract_cmd = env_cmd
        return

    on_path = shutil.which("tesseract")
    if on_path:
        pytesseract.pytesseract.tesseract_cmd = on_path
        return

    common_windows_paths = [
        Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe"),
        Path(r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"),
    ]
    for candidate in common_windows_paths:
        if candidate.exists():
            pytesseract.pytesseract.tesseract_cmd = str(candidate)
            return


def _preprocess_image_sharp_pil(
    image: Image.Image,
    min_text_height_px: int,
    max_dimension_px: int,
) -> Image.Image:
    """PIL-only preprocessing used when OpenCV is unavailable."""
    w, h = image.size
    min_dim = min(w, h)
    if min_dim < min_text_height_px:
        scale = min_text_height_px / min_dim
        new_w = min(int(round(w * scale)), max_dimension_px)
        new_h = min(int(round(h * scale)), max_dimension_px)
        image = image.resize((new_w, new_h), Image.Resampling.LANCZOS)
    gray = ImageOps.grayscale(image)
    gray = ImageOps.autocontrast(gray)
    return gray.filter(ImageFilter.SHARPEN)


def _preprocess_image_sharp_cv2(
    image: Image.Image,
    min_text_height_px: int,
    max_dimension_px: int,
) -> Image.Image:
    """OpenCV pipeline for blurry or low-quality document images."""
    if not _CV2_AVAILABLE:
        return _preprocess_image_sharp_pil(image, min_text_height_px, max_dimension_px)

    arr = np.array(image)
    gray = arr if arr.ndim == 2 else cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
    h, w = gray.shape
    min_dim = min(w, h)
    if min_dim < min_text_height_px:
        scale = min_text_height_px / min_dim
        new_w = min(int(round(w * scale)), max_dimension_px)
        new_h = min(int(round(h * scale)), max_dimension_px)
        gray = cv2.resize(gray, (new_w, new_h), interpolation=cv2.INTER_CUBIC)

    denoised = cv2.fastNlMeansDenoising(gray, h=10, templateWindowSize=7, searchWindowSize=21)
    gaussian = cv2.GaussianBlur(denoised, (0, 0), 2.0)
    sharp = cv2.addWeighted(denoised, 1.5, gaussian, -0.5, 0)
    thresh = cv2.adaptiveThreshold(
        sharp,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        15,
        2,
    )
    return Image.fromarray(thresh)


def _preprocess_image_simple(image: Image.Image) -> Image.Image:
    """Minimal preprocessing for faster but less robust OCR."""
    return ImageOps.autocontrast(ImageOps.grayscale(image))


def _preprocess_image_for_ocr(image: Image.Image) -> Image.Image:
    """Apply OCR preprocessing based on current settings."""
    cfg = _get_ocr_config()
    if not cfg.get("sharp_preprocessing", True):
        return _preprocess_image_simple(image)
    min_px = int(cfg.get("min_text_height_px", _DEFAULT_OCR_CONFIG["min_text_height_px"]))
    max_px = int(cfg.get("max_dimension_px", _DEFAULT_OCR_CONFIG["max_dimension_px"]))
    if _CV2_AVAILABLE:
        return _preprocess_image_sharp_cv2(image, min_px, max_px)
    return _preprocess_image_sharp_pil(image, min_px, max_px)


def _run_ocr_robust(processed: Image.Image, use_multi_psm: bool = True) -> str:
    """Run Tesseract and optionally merge results from multiple page modes."""
    config_base = "--oem 3"
    if use_multi_psm:
        texts: List[str] = []
        for psm in (6, 3):
            try:
                value = pytesseract.image_to_string(
                    processed,
                    config=f"{config_base} --psm {psm}",
                ).strip()
                if value:
                    texts.append(value)
            except Exception:
                continue
        if not texts:
            return ""
        return max(texts, key=len)
    return pytesseract.image_to_string(processed, config=f"{config_base} --psm 6").strip()


def _extract_text_from_image(path: Path) -> str:
    """Run preprocessing and OCR on an image file."""
    _configure_tesseract_cmd()
    image = Image.open(path).convert("RGB")
    processed = _preprocess_image_for_ocr(image)
    cfg = _get_ocr_config()
    use_multi_psm = bool(cfg.get("sharp_preprocessing", True))
    return _run_ocr_robust(processed, use_multi_psm=use_multi_psm)


def _extract_text_from_docx(path: Path) -> str:
    """Extract text from DOCX paragraphs and tables."""
    if Document is None:
        raise RuntimeError(
            "DOCX support requires 'python-docx'. Install dependencies from requirements.txt."
        )

    doc = Document(str(path))
    chunks: List[str] = []
    chunks.extend(paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                chunks.append(" | ".join(values))
    return "\n".join(chunks).strip()


def _extract_text_from_pdf_pypdf(path: Path) -> str:
    """Extract text from PDF pages with pypdf, tolerating malformed pages."""
    if PdfReader is None:
        return ""

    try:
        with redirect_stderr(io.StringIO()), redirect_stdout(io.StringIO()):
            reader = PdfReader(str(path), strict=False)
    except TypeError:
        with redirect_stderr(io.StringIO()), redirect_stdout(io.StringIO()):
            reader = PdfReader(str(path))
    except Exception:
        return ""

    chunks: List[str] = []
    for page in reader.pages:
        try:
            with redirect_stderr(io.StringIO()), redirect_stdout(io.StringIO()):
                text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            chunks.append(text)
    return "\n".join(chunks).strip()


def _extract_text_from_pdf_fitz(path: Path) -> str:
    """Extract embedded text with PyMuPDF when pypdf cannot parse the file."""
    if fitz is None:
        return ""

    tools = getattr(fitz, "TOOLS", None)
    if tools is not None:
        if hasattr(tools, "mupdf_display_errors"):
            tools.mupdf_display_errors(False)
        if hasattr(tools, "mupdf_display_warnings"):
            tools.mupdf_display_warnings(False)

    try:
        with fitz.open(str(path)) as doc:
            return "\n".join(page.get_text("text") for page in doc).strip()
    except Exception:
        return ""


def _extract_text_from_pdf(path: Path) -> str:
    """Extract text from PDF pages using local parsers and OCR fallback."""
    text = _extract_text_from_pdf_pypdf(path)
    if text:
        return text

    text = _extract_text_from_pdf_fitz(path)
    if text:
        return text

    if fitz is None:
        raise RuntimeError(
            "This PDF could not be parsed as text. Install 'pymupdf' for resilient PDF parsing and scanned PDF OCR fallback."
        )

    try:
        ocr_chunks: List[str] = []
        _configure_tesseract_cmd()
        cfg = _get_ocr_config()
        scale = int(cfg.get("pdf_scale", _DEFAULT_OCR_CONFIG["pdf_scale"]))
        scale = max(2, min(scale, 5))
        with fitz.open(str(path)) as doc:
            for page in doc:
                pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale))
                image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                processed = _preprocess_image_for_ocr(image)
                use_multi_psm = bool(cfg.get("sharp_preprocessing", True))
                ocr_chunks.append(_run_ocr_robust(processed, use_multi_psm=use_multi_psm))
        text = "\n".join(ocr_chunks).strip()
        if text:
            return text
        raise RuntimeError("No text could be extracted from this PDF.")
    except pytesseract.TesseractNotFoundError as exc:
        raise RuntimeError(
            "Tesseract OCR is not installed or not in PATH. Install Tesseract locally for offline PDF OCR fallback."
        ) from exc


def read_document_text(path: Path) -> str:
    """Read a supported document path and return extracted text."""
    if not path.exists():
        raise FileNotFoundError(f"Input file not found: {path}")
    if not path.is_file():
        raise ValueError(f"Input path is not a file: {path}")

    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return path.read_text(encoding="utf-8", errors="replace")

    if suffix in DOCX_SUFFIXES:
        try:
            return _extract_text_from_docx(path)
        except RuntimeError:
            raise
        except Exception as exc:
            raise RuntimeError(f"Failed to extract text from DOCX: {exc}") from exc

    if suffix in PDF_SUFFIXES:
        try:
            return _extract_text_from_pdf(path)
        except RuntimeError:
            raise
        except Exception as exc:
            raise RuntimeError(f"Failed to extract text from PDF: {exc}") from exc

    if suffix in IMAGE_SUFFIXES:
        try:
            return _extract_text_from_image(path)
        except pytesseract.TesseractNotFoundError as exc:
            raise RuntimeError(
                "Tesseract OCR is not installed or not in PATH. Install Tesseract locally for offline image text extraction."
            ) from exc

    raise ValueError(
        "Unsupported file type. Use text, DOCX, PDF, or image files "
        "(.txt, .md, .csv, .log, .docx, .pdf, .png, .jpg, .jpeg, .bmp, .tiff, .webp)."
    )



